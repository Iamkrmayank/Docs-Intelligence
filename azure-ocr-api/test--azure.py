#!/usr/bin/env python3
"""
Azure-only: PDF → OCR (Azure Document Intelligence) → Markdown → DOCX
- Uses Azure Document Intelligence (prebuilt-layout) ONLY
- Requests figure crops (output=figures) and embeds them near the page text
- Falls back to local crops from figure boundingRegions when server crops aren't returned
- Optional: embed whole-page snapshots if no figures were found (--embed-page-snapshots-when-no-figures)
- Optional Pandoc path renders LaTeX math ($...$, $$...$$) and tables
- Optional region crops (from your own regions.json) are embedded as images
"""

import os, sys, json, base64, requests, argparse, re, io, shutil, tempfile, time
from typing import Dict, Any, List, Tuple, Optional
from pathlib import Path
from datetime import datetime

from dotenv import load_dotenv
load_dotenv()

# ==== Env helpers ==============================================================
def _env(name: str, default: Optional[str] = None) -> Optional[str]:
    val = os.environ.get(name)
    if val is None:
        return default
    # strip quotes and whitespace
    return val.strip().strip('"').strip("'")

# ==== Azure Document Intelligence (prebuilt-layout) ===========================
AZURE_DI_ENDPOINT    = _env("AZURE_DI_ENDPOINT")
AZURE_DI_KEY         = _env("AZURE_DI_KEY")
# Default to GA 2024-11-30; sanitize comments/quotes if present
AZURE_DI_API_VERSION = ((_env("AZURE_DI_API_VERSION", "2024-11-30") or "2024-11-30").split("#")[0]).strip()
AZURE_DI_MODEL_ID    = "prebuilt-layout"  # change if you prefer another model

# ==== Optional Pandoc (for LaTeX math & tables) ===============================
HAVE_PYPANDOC = True
try:
    import pypandoc  # type: ignore
except Exception:
    HAVE_PYPANDOC = False

def ensure_pandoc_available() -> bool:
    if shutil.which("pandoc"):
        return True
    if HAVE_PYPANDOC:
        try:
            pypandoc.download_pandoc()
            return True
        except Exception:
            return False
    return False

# ==== Imaging / PDF / DOCX fallback ==========================================
from PIL import Image
import fitz  # PyMuPDF

from docx import Document
from docx.shared import Inches, Pt

# ==== Basics =================================================================
def die(msg: str, code: int = 1):
    print(f"[ERR] {msg}", file=sys.stderr); sys.exit(code)

def _strip(s: str) -> str:
    return (s or "").strip()

# -----------------------------------------------------------------------------
# Azure Document Intelligence (prebuilt-layout, async LRO) with figure crops
# -----------------------------------------------------------------------------
def azure_di_analyze_pdf(pdf_bytes: bytes) -> Dict[str, Any]:
    """
    Calls Azure DI prebuilt-layout with PDF bytes.
    Primary path: GA API (2024-11-30+) JSON body with base64Source + output=figures.
    Fallback: if the endpoint rejects JSON (older preview), retry with binary body.
    """
    if not (AZURE_DI_ENDPOINT and AZURE_DI_KEY):
        die("Azure DI endpoint/key missing. Set AZURE_DI_ENDPOINT and AZURE_DI_KEY in .env")

    base_url = AZURE_DI_ENDPOINT.rstrip("/")
    url = (
        base_url
        + f"/documentintelligence/documentModels/{AZURE_DI_MODEL_ID}:analyze"
        + f"?api-version={AZURE_DI_API_VERSION}&output=figures"
    )

    # Try JSON base64 first (GA style)
    headers_json = {
        "Ocp-Apim-Subscription-Key": AZURE_DI_KEY,
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    body_json = {"base64Source": base64.b64encode(pdf_bytes).decode("utf-8")}

    print(f"[INFO] analyze url={url}")
    r = requests.post(url, headers=headers_json, json=body_json, timeout=120)
    print(f"[HTTP submit JSON] {r.status_code}")
    if r.status_code in (200, 201, 202):
        op_loc = r.headers.get("operation-location") or r.headers.get("Operation-Location")
        if not op_loc:
            try:
                data = r.json()
                data["_operationLocation"] = None
                return data
            except Exception:
                die("Azure DI: missing operation-location and no JSON body returned.")
        return _poll_azure(op_loc)

    # If GA fails (e.g., 404/415 in some regions), try preview-style binary
    print("[WARN] JSON analyze failed; trying preview-style binary upload...")
    headers_bin = {
        "Ocp-Apim-Subscription-Key": AZURE_DI_KEY,
        "Content-Type": "application/octet-stream",
        "Accept": "application/json",
    }
    r2 = requests.post(url, headers=headers_bin, data=pdf_bytes, timeout=120)
    print(f"[HTTP submit BIN] {r2.status_code}")
    if r2.status_code not in (200, 201, 202):
        print(r.text[:800])
        print(r2.text[:800])
        r2.raise_for_status()

    op_loc2 = r2.headers.get("operation-location") or r2.headers.get("Operation-Location")
    if not op_loc2:
        try:
            data2 = r2.json()
            data2["_operationLocation"] = None
            return data2
        except Exception:
            die("Azure DI: missing operation-location and no JSON body returned (binary fallback).")
    return _poll_azure(op_loc2)

def _poll_azure(op_loc: str) -> Dict[str, Any]:
    poll_headers = {"Ocp-Apim-Subscription-Key": AZURE_DI_KEY, "Accept": "application/json"}
    deadline = time.time() + 300  # up to 5 minutes
    print(f"[INFO] polling {op_loc}")
    while True:
        pr = requests.get(op_loc, headers=poll_headers, timeout=60)
        if pr.status_code >= 400:
            print(pr.text[:2000]); pr.raise_for_status()
        data = pr.json()
        status = (data.get("status") or data.get("operationState") or "").lower()
        if status in ("succeeded", "success", "completed"):
            data["_operationLocation"] = op_loc  # keep the poll URL we used
            print("[INFO] analyze succeeded")
            return data
        if status in ("failed", "error"):
            die(f"Azure DI analyze failed: {json.dumps(data, ensure_ascii=False)[:1000]}")
        if time.time() > deadline:
            die("Azure DI analyze timed out while polling.")
        time.sleep(2)

def azure_di_normalize_pages(di_json: Dict[str, Any]) -> Dict[str, Any]:
    """
    Normalize Azure DI analyze result to { 'pages': [ {'markdown': '...'}, ... ] }
    Prefer page.lines[].content; fallback to global 'content'.
    """
    container = di_json.get("analyzeResult") or di_json.get("result") or di_json
    pages = container.get("pages") or []

    norm_pages = []
    if isinstance(pages, list) and pages:
        for p in pages:
            lines = p.get("lines") or []
            txt = "\n".join(l.get("content","") for l in lines if isinstance(l, dict) and l.get("content"))
            if not _strip(txt):
                txt = _strip(container.get("content") or "")
            norm_pages.append({"markdown": txt or ""})
    else:
        txt = _strip(container.get("content") or "")
        norm_pages = [{"markdown": txt or ""}]

    return {"pages": norm_pages}

# ==== PDF rendering & crop helpers ===========================================
def render_pdf_page_to_image(pdf_path: Path, page_num: int, dpi: int = 300) -> Tuple[Image.Image, Tuple[float,float]]:
    doc = fitz.open(pdf_path.as_posix())
    try:
        page = doc[page_num-1]
        zoom = dpi / 72.0
        pix  = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        img  = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        return img, (page.rect.width, page.rect.height)  # points (1/72 in)
    finally:
        doc.close()

def clamp(v, lo, hi): return max(lo, min(hi, v))

def bbox_to_pixels_from_units(
    bbox_or_poly: Tuple[float, float, float, float],
    page_unit: str,
    page_w_unit: float,
    page_h_unit: float,
    img_w: int,
    img_h: int,
) -> Tuple[int,int,int,int]:
    """
    bbox_or_poly is (x0,y0,x1,y1) in DI "unit" space (PDF ⇒ inches, image ⇒ pixels).
    """
    x0, y0, x1, y1 = bbox_or_poly
    nx0, ny0 = x0 / page_w_unit, y0 / page_h_unit
    nx1, ny1 = x1 / page_w_unit, y1 / page_h_unit
    X0 = int(round(nx0 * img_w)); Y0 = int(round(ny0 * img_h))
    X1 = int(round(nx1 * img_w)); Y1 = int(round(ny1 * img_h))
    x0p, x1p = sorted([X0, X1]); y0p, y1p = sorted([Y0, Y1])
    x0p = clamp(x0p,0,img_w-1); x1p = clamp(x1p,1,img_w)
    y0p = clamp(y0p,0,img_h-1); y1p = clamp(y1p,1,img_h)
    return x0p, y0p, x1p, y1p

def poly_to_bbox(poly: List[Dict[str, float]]) -> Tuple[float,float,float,float]:
    xs = [p["x"] for p in poly]; ys = [p["y"] for p in poly]
    return min(xs), min(ys), max(xs), max(ys)

def crop_and_save(img: Image.Image, bbox_px: Tuple[int,int,int,int],
                  padding: int, out_dir: Path, stem: str) -> Path:
    x0,y0,x1,y1 = bbox_px
    if padding:
        x0 = max(0, x0-padding); y0 = max(0, y0-padding)
        x1 = min(img.width, x1+padding); y1 = min(img.height, y1+padding)
    crop = img.crop((x0,y0,x1,y1)).convert("RGB")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{stem}.jpg"
    crop.save(out_path, quality=92, optimize=True)
    return out_path

# ==== Markdown helpers (math/tables) ==========================================
def clean_markdown(md: str) -> str:
    md = md.replace("\r\n", "\n").replace("\r", "\n")
    return md

_MATH_OR_TABLE_RE = re.compile(
    r"(\$\$.*?\$\$|\$[^$\n]+\$|\\\(|\\\)|\\\[|\\\]|\\begin\{(equation|align|eqnarray|gather|aligned)\}|(^\s*\|.*\|\s*$\n^\s*\|?\s*[-:]+\s*(\|[-:]+\s*)+$))",
    re.MULTILINE | re.DOTALL
)

def detect_math_or_tables(pages_text: List[str]) -> bool:
    joined = "\n\n".join(pages_text)
    return bool(_MATH_OR_TABLE_RE.search(joined))

_IMG_MD_RE   = re.compile(r'!\[[^\]]*\]\([^)]+\)', re.IGNORECASE)
_IMG_HTML_RE = re.compile(r'<img\b[^>]*>', re.IGNORECASE)

def strip_inline_images(md: str) -> str:
    md = _IMG_MD_RE.sub('', md)
    md = _IMG_HTML_RE.sub('', md)
    return md

def md_image(path: Path, width_in: float) -> str:
    p = Path(path).resolve().as_posix()
    return f'![]({p}){{width={width_in}in}}'

def build_markdown(pages_text: List[str],
                   crops_by_page: Dict[int, List[Path]],
                   insert_page_breaks: bool,
                   image_max_width_in: float) -> str:
    parts: List[str] = []
    for i, txt in enumerate(pages_text, start=1):
        text_clean = strip_inline_images(_strip(txt))
        parts.append(f"\n\n## Page {i}\n\n{text_clean}\n")
        for im in crops_by_page.get(i, []):
            parts.append("\n" + md_image(im, image_max_width_in) + "\n")
        if insert_page_breaks and i < len(pages_text):
            parts.append("\n\\newpage\n")
    return clean_markdown("".join(parts)).strip() + "\n"

def build_docx_with_pandoc_to_path(md_text: str, out_path: Path, resource_dirs: Optional[List[Path]] = None) -> None:
    if not HAVE_PYPANDOC:
        raise RuntimeError("pypandoc not installed")
    if not ensure_pandoc_available():
        raise RuntimeError("Pandoc not available and auto-download failed")
    with tempfile.TemporaryDirectory() as td:
        md_file = Path(td) / "in.md"
        md_file.write_text(md_text, encoding="utf-8")
        extra_args = ["--standalone"]
        if resource_dirs:
            search_path = os.pathsep.join(str(Path(p).resolve()) for p in resource_dirs)
            extra_args.append(f"--resource-path={search_path}")
        pypandoc.convert_file(
            str(md_file),
            to="docx",
            format="gfm+tex_math_dollars+pipe_tables",
            outputfile=str(out_path),
            extra_args=extra_args,
        )

# ==== Basic DOCX fallback (no native math) ====================================
def add_text_block(doc: Document, text: str):
    for chunk in text.split("\n\n"):
        chunk = chunk.strip()
        if chunk:
            doc.add_paragraph(chunk)

def build_docx_with_python_docx_to_path(pages_text: List[str],
                                        crops_by_page: Dict[int, List[Path]],
                                        out_path: Path,
                                        insert_page_breaks: bool,
                                        image_max_width_in: float) -> None:
    doc = Document()
    style = doc.styles['Normal'].font
    style.name = "Calibri"
    style.size = Pt(11)

    for i, txt in enumerate(pages_text, start=1):
        doc.add_heading(f"Page {i}", level=2)
        if _strip(txt):
            add_text_block(doc, txt)
        for im in crops_by_page.get(i, []):
            doc.add_picture(str(im), width=Inches(image_max_width_in))
        if insert_page_breaks and i < len(pages_text):
            doc.add_page_break()
    doc.save(out_path)

# ==== Regions JSON (manual crops) =============================================
def render_pdf_page_meta(pdf_path: Path, page_num: int, dpi: int = 300):
    img, pts_wh = render_pdf_page_to_image(pdf_path, page_num, dpi=dpi)
    return img, pts_wh

def crops_from_regions(pdf_path: Path, regions_json_path: Path, dpi: int, assets_dir: Path) -> Dict[int, List[Path]]:
    with open(regions_json_path, "r", encoding="utf-8") as jf:
        cfg = json.load(jf)
    pages_cfg = cfg.get("pages", {})
    out: Dict[int, List[Path]] = {}
    for p_str, regs in pages_cfg.items():
        try:
            pnum = int(p_str)
        except:
            continue
        if not isinstance(regs, list):
            continue
        page_img, pts_wh = render_pdf_page_meta(pdf_path, pnum, dpi=dpi)
        for k, r in enumerate(regs, start=1):
            coords     = r["coords"]
            coord_type = r.get("coord_type","norm")
            padding    = int(r.get("padding", 8))
            label      = r.get("label", f"crop{k}")

            if coord_type == "norm":
                x0,y0,x1,y1 = coords
                bbox_px = (
                    int(round(x0*page_img.width)),
                    int(round(y0*page_img.height)),
                    int(round(x1*page_img.width)),
                    int(round(y1*page_img.height)),
                )
            elif coord_type == "pdf_points":
                pts_w, pts_h = pts_wh
                bbox_px = bbox_to_pixels_from_units(
                    (coords[0], coords[1], coords[2], coords[3]),
                    page_unit="point",
                    page_w_unit=pts_w, page_h_unit=pts_h,
                    img_w=page_img.width, img_h=page_img.height
                )
            elif coord_type == "pixel":
                bbox_px = (int(coords[0]), int(coords[1]), int(coords[2]), int(coords[3]))
            else:
                raise ValueError("coord_type must be norm|pixel|pdf_points")
            saved = crop_and_save(page_img, bbox_px, padding, assets_dir, f"p{pnum}_{k}_{label}")
            out.setdefault(pnum, []).append(saved)
    return out

# ==== FIGURE crops from Azure DI ==============================================
def extract_figure_crops(
    di_json: Dict[str, Any],
    pdf_path: Path,
    assets_dir: Path,
    dpi: int = 300
) -> Dict[int, List[Path]]:
    """
    Prefer server-provided figure crops (when output=figures) and fall back to local cropping
    using boundingRegions → polygon → bbox on a rendered page image.
    Returns: { page_number: [Path, ...], ... }
    """
    container = di_json.get("analyzeResult") or di_json.get("result") or di_json
    figures = container.get("figures") or []
    pages_meta = container.get("pages") or []  # has page width/height + unit
    op_loc = di_json.get("_operationLocation")

    out: Dict[int, List[Path]] = {}

    def try_download_server_figure(fig_id: str) -> Optional[bytes]:
        if not op_loc:
            return None
        try:
            # Support both GA and preview poll URLs
            if "/analyzeResults/" in op_loc:
                base_results_url = op_loc.split("?")[0]  # .../analyzeResults/{id}
                url = f"{base_results_url}/figures/{fig_id}?api-version={AZURE_DI_API_VERSION}"
            elif "/operations/" in op_loc:
                base = op_loc.split("/operations/")[0].rstrip("/")
                result_id = op_loc.split("/operations/")[1].split("?")[0]
                url = f"{base}/analyzeResults/{result_id}/figures/{fig_id}?api-version={AZURE_DI_API_VERSION}"
            else:
                return None

            for accept in ("image/png", "image/jpeg"):
                h = {"Ocp-Apim-Subscription-Key": AZURE_DI_KEY, "Accept": accept}
                r = requests.get(url, headers=h, timeout=60)
                if r.status_code == 200 and r.content:
                    return r.content
            return None
        except Exception:
            return None

    # cache rendered pages for local fallback
    rendered_cache: Dict[int, Tuple[Image.Image, Tuple[float,float]]] = {}

    for idx, fig in enumerate(figures, start=1):
        brs = fig.get("boundingRegions") or []
        if not brs:
            continue
        page_no = int(brs[0].get("pageNumber", 1))
        polygon = brs[0].get("polygon") or []

        if len(polygon) >= 4 and isinstance(polygon[0], dict):
            bbox_unit = poly_to_bbox(polygon)
        else:
            xs = polygon[0::2]; ys = polygon[1::2]
            bbox_unit = (min(xs), min(ys), max(xs), max(ys))

        saved_path: Optional[Path] = None

        # server crop first
        fig_id = fig.get("id")
        if fig_id:
            img_bytes = try_download_server_figure(fig_id)
            if img_bytes:
                saved_path = assets_dir / f"p{page_no}_figure_{idx}.jpg"
                saved_path.write_bytes(img_bytes)

        # local crop fallback
        if not saved_path:
            if page_no not in rendered_cache:
                rendered_cache[page_no] = render_pdf_page_to_image(pdf_path, page_no, dpi=dpi)
            page_img, pts_wh = rendered_cache[page_no]
            page_meta = next((p for p in pages_meta if int(p.get("pageNumber", 0)) == page_no), None)
            if not page_meta:
                pts_w, pts_h = pts_wh
                in_w = pts_w / 72.0
                in_h = pts_h / 72.0
                bbox_px = bbox_to_pixels_from_units(
                    bbox_unit, page_unit="inch",
                    page_w_unit=in_w, page_h_unit=in_h,
                    img_w=page_img.width, img_h=page_img.height
                )
            else:
                unit = (page_meta.get("unit") or "").lower()  # "inch" for PDFs; "pixel" for images
                pw = float(page_meta.get("width", 0.0))
                ph = float(page_meta.get("height", 0.0))
                bbox_px = bbox_to_pixels_from_units(
                    bbox_unit, page_unit=unit, page_w_unit=pw, page_h_unit=ph,
                    img_w=page_img.width, img_h=page_img.height
                )

            saved_path = crop_and_save(page_img, bbox_px, padding=8, out_dir=assets_dir,
                                       stem=f"p{page_no}_figure_{idx}")

        out.setdefault(page_no, []).append(saved_path)

    return out

# ==== Page snapshot helper =====================================================
def render_page_snapshot(
    pdf_path: Path, page_no: int, assets_dir: Path, dpi: int, stem_prefix: str = "page"
) -> Path:
    img, _ = render_pdf_page_to_image(pdf_path, page_no, dpi=dpi)
    out = assets_dir / f"{stem_prefix}_{page_no}.jpg"
    img.save(out, quality=88, optimize=True)
    return out

# ==== Main ====================================================================
def main():
    ap = argparse.ArgumentParser(description="OCR PDF → DOCX (Azure Document Intelligence only). Pandoc path for math/tables; optional figure crops and fallbacks.")
    ap.add_argument("pdf", nargs="+", help="Path to the PDF (quote if it has spaces)")
    ap.add_argument("--title", default=None, help="DOCX title heading (markdown H1)")
    ap.add_argument("--out", default=None, help="Base output name (no ext)")
    ap.add_argument("--regions-json", default=None, help="Crop & embed regions from JSON")
    ap.add_argument("--dpi", type=int, default=300, help="Render DPI for cropping (regions/figures fallback)")
    ap.add_argument("--image-max-width", type=float, default=6.5, help="Image width in inches")
    ap.add_argument("--no-page-breaks", dest="no_page_breaks", action="store_true", help="No page breaks between pages")
    ap.add_argument("--prefer-pandoc", action="store_true", help="Prefer Pandoc if available (recommended)")
    ap.add_argument("--force-pandoc", action="store_true", help="Force Pandoc; error if unavailable")

    # New/compat flags for images
    ap.add_argument("--include-images", action="store_true", default=True,
                    help="(compat) Include DI-detected figures. Default: on")
    ap.add_argument("--no-figures", action="store_true",
                    help="Disable inserting DI figures/crops")

    # Optional: add whole-page snapshot if a page has zero figures
    ap.add_argument("--embed-page-snapshots-when-no-figures", action="store_true",
                    help="Embed a page snapshot image for pages where no figures were found.")
    ap.add_argument("--snapshot-dpi", type=int, default=220,
                    help="DPI used for page snapshots when --embed-page-snapshots-when-no-figures is set.")

    args = ap.parse_args()

    path = " ".join(args.pdf)
    pdf_path = Path(path)
    if not pdf_path.exists():
        die(f"File not found: {pdf_path}")

    pdf_bytes = pdf_path.read_bytes()
    print(f"[INFO] file={pdf_path.name} bytes={len(pdf_bytes)}")
    print(f"[INFO] api-version={AZURE_DI_API_VERSION}")

    # OCR with Azure DI (with output=figures)
    ocr_json = azure_di_analyze_pdf(pdf_bytes)
    Path("ocr_response.json").write_text(json.dumps(ocr_json, ensure_ascii=False, indent=2), encoding="utf-8")
    print("[INFO] wrote ocr_response.json")

    container = azure_di_normalize_pages(ocr_json)
    pages = container.get("pages")
    if not isinstance(pages, list) or not pages:
        die("No pages found in Azure DI response.")

    # Resolve output dir + stem safely (supports --out as "name" OR "/path/name")
    if args.out:
        out_base = Path(args.out)
        if out_base.is_absolute() or out_base.parent != Path():
            # --out includes a directory or absolute path
            out_dir = out_base.parent
            out_stem = out_base.stem
        else:
            # --out is just a bare name
            out_dir = pdf_path.parent
            out_stem = out_base.name
    else:
        out_dir = pdf_path.parent
        out_stem = pdf_path.stem

    docx_path = out_dir / f"{out_stem}_ocr.docx"
    assets_dir = out_dir / f"{out_stem}_assets"


    if assets_dir.exists():
        shutil.rmtree(assets_dir, ignore_errors=True)
    assets_dir.mkdir(parents=True, exist_ok=True)

    # Gather page texts
    pages_text: List[str] = []
    print(f"[OK] pages={len(pages)}")
    for i, p in enumerate(pages, start=1):
        txt = p.get("markdown") or ""
        if args.title and i == 1:
            txt = f"# {args.title}\n\n{txt}"
        pages_text.append(txt)

    # FIGURE crops (preferred) — gate with flags
    crops_by_page: Dict[int, List[Path]] = {}
    if not args.no_figures and args.include_images:
        crops_by_page = extract_figure_crops(
            ocr_json, pdf_path, assets_dir, dpi=args.dpi
        )
        total_figs = sum(len(v) for v in crops_by_page.values())
        print(f"[INFO] embedded {total_figs} figure crop(s) from Azure DI")
    else:
        print("[INFO] figure insertion disabled by flags")

    # Optional: page snapshots where a page has no figures
    if args.embed_page_snapshots_when_no_figures:
        container_full = (ocr_json.get("analyzeResult") or ocr_json.get("result") or ocr_json)
        di_pages = container_full.get("pages") or []
        total_pages = max((int(p.get("pageNumber", 0)) for p in di_pages), default=len(pages_text))
        for page_no in range(1, total_pages + 1):
            if len(crops_by_page.get(page_no, [])) == 0:
                snap_path = render_page_snapshot(pdf_path, page_no, assets_dir, dpi=args.snapshot_dpi, stem_prefix="snapshot")
                crops_by_page.setdefault(page_no, []).append(snap_path)
        print("[INFO] page snapshots embedded where no figures were found")

    # Optional extra crops from regions.json → merge
    if args.regions_json:
        rj = Path(args.regions_json)
        if not rj.exists():
            die(f"regions JSON not found: {rj}")
        extra = crops_from_regions(pdf_path, rj, args.dpi, assets_dir)
        for k, v in extra.items():
            crops_by_page.setdefault(k, []).extend(v)
        print(f"[INFO] merged {sum(len(v) for v in extra.values())} manual crop(s) from regions")

    # Decide pipeline (Pandoc vs python-docx)
    wants_math = detect_math_or_tables(pages_text)
    prefer_pd  = args.prefer_pandoc or args.force_pandoc or wants_math

    try:
        if prefer_pd:
            if not ensure_pandoc_available():
                if args.force_pandoc:
                    die("Pandoc required (--force-pandoc) but not available.")
                else:
                    print("[WARN] Pandoc unavailable; falling back to basic DOCX (math will be plain text).")
                    raise RuntimeError("no-pandoc")
            md_text = build_markdown(
                pages_text,
                crops_by_page=crops_by_page,
                insert_page_breaks=not args.no_page_breaks,
                image_max_width_in=args.image_max_width
            )
            build_docx_with_pandoc_to_path(md_text, docx_path, resource_dirs=[assets_dir])
            print(f"[OK] wrote docx (pandoc): {docx_path}")
        else:
            raise RuntimeError("prefer-basic")
    except Exception:
        # Fallback: python-docx (no native equations)
        build_docx_with_python_docx_to_path(
            pages_text,
            crops_by_page=crops_by_page,
            out_path=docx_path,
            insert_page_breaks=not args.no_page_breaks,
            image_max_width_in=args.image_max_width
        )
        print(f"[OK] wrote docx (basic): {docx_path}")

if __name__ == "__main__":
    main()
